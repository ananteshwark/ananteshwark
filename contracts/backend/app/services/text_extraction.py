"""Extract text from PDFs (digital + scanned via OCR), DOCX, and images.

OCR also records where each word sat on the page. A scanned contract has no
text layer, so the PDF viewer has nothing to shade and risk highlighting simply
does not work on scans — the one case where the contract is *only* an image and
the reader most needs the flags drawn for them. Capturing per-word boxes at OCR
time gives the overlay real coordinates to draw against.

The text and the boxes are built from a single `image_to_data` call rather than
`image_to_string` for the text and a second pass for the boxes. Two passes can
disagree, and a word present in the stored text but absent from the stored
boxes is a flag that silently fails to place.
"""
import logging
from pathlib import Path

log = logging.getLogger(__name__)

# Under this many characters, a PDF page is assumed to be scanned and OCR is tried
MIN_TEXT_PER_PAGE = 50

# Rendering DPI for scanned pages. Word boxes are in pixels at this DPI, and the
# viewer rescales them to whatever width it draws the page at, so the value only
# has to be recorded, not agreed on.
OCR_DPI = 300

# Words below this Tesseract confidence are noise — stray specks and scan edges
# that would otherwise litter the page with tiny highlight boxes.
_MIN_WORD_CONF = 30

# A hard ceiling on stored boxes per document. A 300-page scan at ~600 words a
# page would otherwise put ~180k objects in one JSON column; past this the
# layout is dropped rather than bloating every read of the row.
_MAX_WORDS = 120_000


class TextExtractionError(Exception):
    pass


def _ocr_page(image) -> tuple[str, dict]:
    """OCR one page image. Returns (text, {"w","h","words"}).

    Each word is {"t","x","y","w","h"} in pixels of this image. Keys are short
    because they repeat once per word and this is stored as JSON.
    """
    try:
        import pytesseract
        from pytesseract import Output, TesseractNotFoundError
    except ImportError as exc:
        raise TextExtractionError(
            "pytesseract is not installed — run `pip install pytesseract` (and install the "
            "'tesseract-ocr' system package) to OCR scanned documents"
        ) from exc
    # Grayscale improves OCR accuracy on scanned pages.
    try:
        image = image.convert("L")
    except Exception:  # pragma: no cover - non-fatal preprocessing
        # OCR still runs, just on the colour image and less accurately. Worth a
        # line in the log when a page's extraction comes out poorly.
        log.warning("Could not convert page to grayscale for OCR; continuing "
                    "with the original image", exc_info=True)
    try:
        data = pytesseract.image_to_data(image, output_type=Output.DICT)
    except TesseractNotFoundError as exc:
        raise TextExtractionError(
            "The Tesseract OCR engine is not installed — install the 'tesseract-ocr' system "
            "package (e.g. `sudo apt install tesseract-ocr`) to read scanned documents"
        ) from exc

    words: list[dict] = []
    lines: list[str] = []
    current: list[str] = []
    key = None
    for i in range(len(data.get("text") or [])):
        raw = (data["text"][i] or "").strip()
        try:
            conf = float(data["conf"][i])
        except (TypeError, ValueError):
            conf = -1.0
        # Tesseract emits a row per layout element, not just per word; the
        # non-word rows carry empty text. Line grouping still has to advance on
        # them, or two paragraphs run together.
        here = (data.get("block_num", [0] * (i + 1))[i],
                data.get("par_num", [0] * (i + 1))[i],
                data.get("line_num", [0] * (i + 1))[i])
        if key is not None and here != key and current:
            lines.append(" ".join(current))
            current = []
        key = here
        if not raw or conf < _MIN_WORD_CONF:
            continue
        current.append(raw)
        words.append({
            "t": raw,
            "x": int(data["left"][i]), "y": int(data["top"][i]),
            "w": int(data["width"][i]), "h": int(data["height"][i]),
        })
    if current:
        lines.append(" ".join(current))

    size = getattr(image, "size", None) or (0, 0)
    return "\n".join(lines), {"w": int(size[0]), "h": int(size[1]), "words": words}


def _ocr_image(image) -> str:
    return _ocr_page(image)[0]


def _extract_pdf(path: Path) -> tuple[str, dict | None]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n".join(pages)

    # Enough of a digital text layer -> use it directly. No layout is recorded:
    # the viewer reads coordinates straight off the page in that case, and a
    # second set from OCR would only be a chance to disagree with it.
    if len(text.strip()) >= MIN_TEXT_PER_PAGE * max(len(pages), 1):
        return text, None

    # Otherwise treat as scanned and OCR every page. Surface a clear, actionable
    # error when an OCR dependency is missing (the usual cause of "can't extract
    # from scanned documents") instead of silently passing empty text to the AI.
    try:
        from pdf2image import convert_from_path
    except ImportError as exc:
        raise TextExtractionError(
            "pdf2image is not installed — required to OCR scanned PDFs "
            "(`pip install pdf2image`, plus the 'poppler-utils' system package)"
        ) from exc

    try:
        images = convert_from_path(str(path), dpi=OCR_DPI)
    except TextExtractionError:
        raise
    except Exception as exc:  # poppler missing, corrupt file, ...
        detail = str(exc).lower()
        if any(k in detail for k in ("poppler", "pdfinfo", "pdftoppm", "pdftocairo")):
            raise TextExtractionError(
                "poppler-utils is not installed — required to render scanned PDFs for OCR "
                "(e.g. `sudo apt install poppler-utils`)"
            ) from exc
        if not text.strip():
            raise TextExtractionError(f"Could not render the scanned PDF for OCR: {exc}") from exc
        return text, None  # keep whatever little text we did have

    ocr_pages, page_texts = [], []
    for img in images:
        page_text, layout = _ocr_page(img)
        page_texts.append(page_text)
        ocr_pages.append(layout)
    ocr_text = "\n".join(page_texts)

    # Keep the digital text if OCR read less than it did — and drop the layout
    # with it, since boxes that describe text nobody is searching are worse than
    # none: they would place highlights against words the reader never sees.
    if len(ocr_text.strip()) <= len(text.strip()):
        return text, None

    total = sum(len(p["words"]) for p in ocr_pages)
    if total > _MAX_WORDS:
        log.warning("OCR layout for %s has %d words (limit %d) — storing text only; "
                    "risk highlighting will fall back to the flag list on this document",
                    path.name, total, _MAX_WORDS)
        return ocr_text, None
    return ocr_text, {"dpi": OCR_DPI, "pages": ocr_pages}


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_image(path: Path) -> tuple[str, dict | None]:
    from PIL import Image

    with Image.open(str(path)) as img:
        text, layout = _ocr_page(img)
    if not layout["words"]:
        return text, None
    return text, {"dpi": OCR_DPI, "pages": [layout]}


def extract_text_with_layout(path: str | Path) -> tuple[str, dict | None]:
    """Extract text and, for OCR'd documents, where each word sits on the page.

    The layout is None for anything with a real text layer (digital PDFs, DOCX)
    — the viewer gets coordinates from the document itself in those cases.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, layout = _extract_pdf(path)
    elif suffix == ".docx":
        text, layout = _extract_docx(path), None
    elif suffix in (".jpg", ".jpeg", ".png"):
        text, layout = _extract_image(path)
    else:
        raise TextExtractionError(f"Unsupported file type: {suffix}")

    if not text.strip():
        raise TextExtractionError("No text could be extracted from the document")
    return text, layout


def extract_text(path: str | Path) -> str:
    """Text only. Callers that have no use for coordinates keep this signature."""
    return extract_text_with_layout(path)[0]
