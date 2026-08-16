"""Extract text from PDFs (digital + scanned via OCR), DOCX, and images."""
import logging
from pathlib import Path

log = logging.getLogger(__name__)

# Under this many characters, a PDF page is assumed to be scanned and OCR is tried
MIN_TEXT_PER_PAGE = 50


class TextExtractionError(Exception):
    pass


def _ocr_image(image) -> str:
    try:
        import pytesseract
        from pytesseract import TesseractNotFoundError
    except ImportError as exc:
        raise TextExtractionError(
            "pytesseract is not installed — run `pip install pytesseract` (and install the "
            "'tesseract-ocr' system package) to OCR scanned documents"
        ) from exc
    # Grayscale improves OCR accuracy on scanned pages.
    try:
        image = image.convert("L")
    except Exception:  # pragma: no cover - non-fatal preprocessing
        pass
    try:
        return pytesseract.image_to_string(image)
    except TesseractNotFoundError as exc:
        raise TextExtractionError(
            "The Tesseract OCR engine is not installed — install the 'tesseract-ocr' system "
            "package (e.g. `sudo apt install tesseract-ocr`) to read scanned documents"
        ) from exc


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n".join(pages)

    # Enough of a digital text layer -> use it directly.
    if len(text.strip()) >= MIN_TEXT_PER_PAGE * max(len(pages), 1):
        return text

    # Otherwise treat as scanned and OCR every page. Surface a clear, actionable
    # error when an OCR dependency is missing (the usual cause of "can't extract
    # from scanned documents") instead of silently passing empty text to the AI.
    try:
        from pdf2image import convert_from_path
    except ImportError as exc:
        raise TextExtractionError(
            "pdf2image is not installed — required to OCR scanned PDFs "
            "(`pip install pdf2image`, plus the 'poppler-utils' system package)"
        ) from exc

    try:
        images = convert_from_path(str(path), dpi=300)
    except TextExtractionError:
        raise
    except Exception as exc:  # poppler missing, corrupt file, ...
        detail = str(exc).lower()
        if any(k in detail for k in ("poppler", "pdfinfo", "pdftoppm", "pdftocairo")):
            raise TextExtractionError(
                "poppler-utils is not installed — required to render scanned PDFs for OCR "
                "(e.g. `sudo apt install poppler-utils`)"
            ) from exc
        if not text.strip():
            raise TextExtractionError(f"Could not render the scanned PDF for OCR: {exc}") from exc
        return text  # keep whatever little text we did have

    ocr_text = "\n".join(_ocr_image(img) for img in images)
    return ocr_text if len(ocr_text.strip()) > len(text.strip()) else text


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_image(path: Path) -> str:
    from PIL import Image

    with Image.open(str(path)) as img:
        return _ocr_image(img)


def extract_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix in (".jpg", ".jpeg", ".png"):
        text = _extract_image(path)
    else:
        raise TextExtractionError(f"Unsupported file type: {suffix}")

    if not text.strip():
        raise TextExtractionError("No text could be extracted from the document")
    return text
